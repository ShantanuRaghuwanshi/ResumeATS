from docx import Document


def create_modern_template(path):
    doc = Document()
    doc.add_heading("{{ personal_details.name }}", 0)
    doc.add_paragraph("Contact: {{ personal_details.contact }}")
    doc.add_paragraph("LinkedIn: {{ personal_details.linkedin }}")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("{{ education }}")
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("{{ work_experience }}")
    doc.add_heading("Projects", level=1)
    doc.add_paragraph("{{ projects }}")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("{{ skills }}")
    doc.save(path)


if __name__ == "__main__":
    create_modern_template("templates/modern.docx")
