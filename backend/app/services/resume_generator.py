from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os
import json
import zipfile
from typing import Dict, List, Optional, Union, Any, Literal
from datetime import datetime
from uuid import uuid4
from jinja2 import Template
from pathlib import Path
import asyncio
import aiofiles


class ResumeExportFormat:
    """Supported export formats"""

    DOCX = "docx"
    PDF = "pdf"
    TXT = "txt"
    JSON = "json"
    HTML = "html"


class ResumeTemplate:
    """Resume template configuration"""

    def __init__(self, name: str, path: str, customizations: Dict[str, Any] = None):
        self.name = name
        self.path = path
        self.customizations = customizations or {}


class ExportOptions:
    """Export configuration options"""

    def __init__(
        self,
        format: str = ResumeExportFormat.DOCX,
        template: str = "modern",
        filename: Optional[str] = None,
        include_metadata: bool = True,
        ats_optimized: bool = False,
        custom_styling: Dict[str, Any] = None,
        sections_to_include: List[str] = None,
    ):
        self.format = format
        self.template = template
        self.filename = filename or f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        self.include_metadata = include_metadata
        self.ats_optimized = ats_optimized
        self.custom_styling = custom_styling or {}
        self.sections_to_include = sections_to_include or []


class BatchExportRequest:
    """Batch export configuration"""

    def __init__(
        self,
        versions: List[Dict[str, Any]],
        export_options: List[ExportOptions],
        output_format: Literal["zip", "individual"] = "zip",
        batch_name: Optional[str] = None,
    ):
        self.versions = versions
        self.export_options = export_options
        self.output_format = output_format
        self.batch_name = (
            batch_name or f"resume_batch_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )


class EnhancedResumeGenerator:
    """Enhanced resume generator with multiple format support and customization"""

    def __init__(self, templates_dir: str = "templates", output_dir: str = "generated"):
        self.templates_dir = Path(templates_dir)
        self.output_dir = Path(output_dir)
        self.templates_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)

        # Available templates with enhanced customization options
        self.templates = {
            "modern": ResumeTemplate(
                "modern",
                "modern.docx",
                customizations={
                    "section_order": [
                        "personal_details",
                        "summary",
                        "work_experience",
                        "skills",
                        "education",
                        "projects",
                    ],
                    "color_scheme": "blue",
                    "font_family": "Calibri",
                    "emphasis_style": "bold",
                    "spacing": "standard",
                },
            ),
            "classic": ResumeTemplate(
                "classic",
                "classic.docx",
                customizations={
                    "section_order": [
                        "personal_details",
                        "work_experience",
                        "education",
                        "skills",
                        "projects",
                    ],
                    "color_scheme": "black",
                    "font_family": "Times New Roman",
                    "emphasis_style": "underline",
                    "spacing": "compact",
                },
            ),
            "minimal": ResumeTemplate(
                "minimal",
                "minimal.docx",
                customizations={
                    "section_order": [
                        "personal_details",
                        "work_experience",
                        "education",
                        "skills",
                    ],
                    "color_scheme": "gray",
                    "font_family": "Arial",
                    "emphasis_style": "bold",
                    "spacing": "wide",
                },
            ),
            "creative": ResumeTemplate(
                "creative",
                "creative.docx",
                customizations={
                    "section_order": [
                        "personal_details",
                        "summary",
                        "skills",
                        "work_experience",
                        "projects",
                        "education",
                    ],
                    "color_scheme": "teal",
                    "font_family": "Helvetica",
                    "emphasis_style": "color",
                    "spacing": "standard",
                },
            ),
            "ats_friendly": ResumeTemplate(
                "ats_friendly",
                "ats_friendly.docx",
                customizations={
                    "section_order": [
                        "personal_details",
                        "summary",
                        "work_experience",
                        "education",
                        "skills",
                        "projects",
                    ],
                    "color_scheme": "none",
                    "font_family": "Arial",
                    "emphasis_style": "caps",
                    "spacing": "standard",
                    "ats_optimized": True,
                },
            ),
        }

    async def generate_resume(
        self,
        resume_data: Dict[str, Any],
        export_options: ExportOptions,
        optimization_results: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Generate a single resume with specified options"""

        # Apply optimization results if provided
        if optimization_results:
            resume_data = self._apply_optimization_results(
                resume_data, optimization_results
            )

        # Apply template customizations
        customized_data = self._apply_template_customizations(
            resume_data, export_options
        )

        # Generate based on format
        if export_options.format == ResumeExportFormat.DOCX:
            return await self._generate_docx(customized_data, export_options)
        elif export_options.format == ResumeExportFormat.PDF:
            return await self._generate_pdf(customized_data, export_options)
        elif export_options.format == ResumeExportFormat.TXT:
            return await self._generate_txt(customized_data, export_options)
        elif export_options.format == ResumeExportFormat.JSON:
            return await self._generate_json(customized_data, export_options)
        elif export_options.format == ResumeExportFormat.HTML:
            return await self._generate_html(customized_data, export_options)
        else:
            raise ValueError(f"Unsupported format: {export_options.format}")

    async def batch_export(self, batch_request: BatchExportRequest) -> str:
        """Export multiple resume versions in batch"""

        if batch_request.output_format == "zip":
            return await self._batch_export_zip(batch_request)
        else:
            return await self._batch_export_individual(batch_request)

    def _apply_optimization_results(
        self, resume_data: Dict[str, Any], optimization_results: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Apply optimization results to resume data"""

        optimized_data = resume_data.copy()

        # Apply section-specific optimizations
        if "section_optimizations" in optimization_results:
            for section, optimizations in optimization_results[
                "section_optimizations"
            ].items():
                if section in optimized_data:
                    # Apply keyword optimizations
                    if "keywords" in optimizations:
                        optimized_data[section] = self._apply_keyword_optimizations(
                            optimized_data[section], optimizations["keywords"]
                        )

                    # Apply content improvements
                    if "content_improvements" in optimizations:
                        optimized_data[section] = self._apply_content_improvements(
                            optimized_data[section],
                            optimizations["content_improvements"],
                        )

                    # Apply suggested text replacements
                    if "text_replacements" in optimizations:
                        optimized_data[section] = self._apply_text_replacements(
                            optimized_data[section], optimizations["text_replacements"]
                        )

        # Apply ATS optimizations
        if "ats_optimizations" in optimization_results:
            optimized_data = self._apply_ats_optimizations(
                optimized_data, optimization_results["ats_optimizations"]
            )

        # Apply job-specific optimizations
        if "job_specific_optimizations" in optimization_results:
            optimized_data = self._apply_job_specific_optimizations(
                optimized_data, optimization_results["job_specific_optimizations"]
            )

        return optimized_data

    def _apply_template_customizations(
        self, resume_data: Dict[str, Any], export_options: ExportOptions
    ) -> Dict[str, Any]:
        """Apply template-specific customizations"""

        customized_data = resume_data.copy()
        template = self.templates.get(export_options.template)

        if template and template.customizations:
            # Apply template-specific formatting
            if "section_order" in template.customizations:
                customized_data = self._reorder_sections(
                    customized_data, template.customizations["section_order"]
                )

            # Apply custom styling
            if export_options.custom_styling:
                customized_data["_styling"] = export_options.custom_styling

        # Filter sections if specified
        if export_options.sections_to_include:
            filtered_data = {}
            for section in export_options.sections_to_include:
                if section in customized_data:
                    filtered_data[section] = customized_data[section]
            # Always include personal details
            if "personal_details" in customized_data:
                filtered_data["personal_details"] = customized_data["personal_details"]
            customized_data = filtered_data

        return customized_data

    async def _generate_docx(
        self, resume_data: Dict[str, Any], export_options: ExportOptions
    ) -> str:
        """Generate DOCX format resume"""

        template_path = self.templates_dir / f"{export_options.template}.docx"
        output_path = self.output_dir / f"{export_options.filename}.docx"

        if template_path.exists():
            # Use template
            doc = DocxTemplate(str(template_path))
            doc.render(resume_data)
            doc.save(str(output_path))
        else:
            # Create from scratch
            doc = Document()
            await self._build_docx_from_scratch(doc, resume_data, export_options)
            doc.save(str(output_path))

        return str(output_path)

    async def _generate_pdf(
        self, resume_data: Dict[str, Any], export_options: ExportOptions
    ) -> str:
        """Generate PDF format resume"""

        # First generate DOCX, then convert to PDF
        docx_options = ExportOptions(
            format=ResumeExportFormat.DOCX,
            template=export_options.template,
            filename=f"{export_options.filename}_temp",
            ats_optimized=export_options.ats_optimized,
            custom_styling=export_options.custom_styling,
        )

        docx_path = await self._generate_docx(resume_data, docx_options)
        pdf_path = self.output_dir / f"{export_options.filename}.pdf"

        try:
            # Try to convert DOCX to PDF using python-docx2pdf
            from docx2pdf import convert

            convert(docx_path, str(pdf_path))

            # Clean up temporary DOCX file
            if os.path.exists(docx_path):
                os.remove(docx_path)

        except ImportError:
            # If docx2pdf is not available, try alternative methods
            try:
                # Try using win32com (Windows only)
                import win32com.client

                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False

                doc = word.Documents.Open(os.path.abspath(docx_path))
                doc.SaveAs(
                    os.path.abspath(str(pdf_path)), FileFormat=17
                )  # 17 = PDF format
                doc.Close()
                word.Quit()

                # Clean up temporary DOCX file
                if os.path.exists(docx_path):
                    os.remove(docx_path)

            except ImportError:
                # If no PDF conversion is available, generate HTML and note the limitation
                html_options = ExportOptions(
                    format=ResumeExportFormat.HTML,
                    template=export_options.template,
                    filename=export_options.filename,
                    ats_optimized=export_options.ats_optimized,
                    custom_styling=export_options.custom_styling,
                )

                html_path = await self._generate_html(resume_data, html_options)

                # Create a note file about PDF conversion
                note_path = self.output_dir / f"{export_options.filename}_pdf_note.txt"
                async with aiofiles.open(note_path, "w", encoding="utf-8") as f:
                    await f.write(
                        "PDF conversion not available. Please install 'docx2pdf' package or use Microsoft Word.\n"
                        f"Generated HTML version available at: {html_path}\n"
                        f"Generated DOCX version available at: {docx_path}"
                    )

                return str(note_path)

        return str(pdf_path)

    async def _generate_txt(
        self, resume_data: Dict[str, Any], export_options: ExportOptions
    ) -> str:
        """Generate TXT format resume (ATS-friendly)"""

        output_path = self.output_dir / f"{export_options.filename}.txt"

        content = []

        # Personal details
        if "personal_details" in resume_data:
            personal = resume_data["personal_details"]
            content.append(personal.get("name", ""))
            if personal.get("email"):
                content.append(f"Email: {personal['email']}")
            if personal.get("phone"):
                content.append(f"Phone: {personal['phone']}")
            if personal.get("linkedin"):
                content.append(f"LinkedIn: {personal['linkedin']}")
            content.append("")

        # Summary/Objective
        if "personal_details" in resume_data and resume_data["personal_details"].get(
            "summary"
        ):
            content.append("SUMMARY")
            content.append("-" * 50)
            content.append(resume_data["personal_details"]["summary"])
            content.append("")

        # Work Experience
        if "work_experience" in resume_data:
            content.append("WORK EXPERIENCE")
            content.append("-" * 50)
            for exp in resume_data["work_experience"]:
                content.append(f"{exp.get('title', '')} - {exp.get('company', '')}")
                if exp.get("from_date") or exp.get("to_date"):
                    content.append(
                        f"{exp.get('from_date', '')} - {exp.get('to_date', 'Present')}"
                    )
                if exp.get("summary"):
                    content.append(exp["summary"])
                if exp.get("achievements"):
                    for achievement in exp["achievements"]:
                        content.append(f"• {achievement}")
                content.append("")

        # Education
        if "education" in resume_data:
            content.append("EDUCATION")
            content.append("-" * 50)
            for edu in resume_data["education"]:
                content.append(f"{edu.get('degree', '')} - {edu.get('university', '')}")
                if edu.get("from_year") or edu.get("to_year"):
                    content.append(
                        f"{edu.get('from_year', '')} - {edu.get('to_year', '')}"
                    )
                content.append("")

        # Skills
        if "skills" in resume_data:
            content.append("SKILLS")
            content.append("-" * 50)
            for skill_cat in resume_data["skills"]:
                content.append(
                    f"{skill_cat.get('category', '')}: {', '.join(skill_cat.get('skills', []))}"
                )
            content.append("")

        # Projects
        if "projects" in resume_data:
            content.append("PROJECTS")
            content.append("-" * 50)
            for project in resume_data["projects"]:
                content.append(project.get("name", ""))
                if project.get("summary"):
                    content.append(project["summary"])
                if project.get("bullets"):
                    for bullet in project["bullets"]:
                        content.append(f"• {bullet}")
                content.append("")

        async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
            await f.write("\n".join(content))

        return str(output_path)

    async def _generate_json(
        self, resume_data: Dict[str, Any], export_options: ExportOptions
    ) -> str:
        """Generate JSON format resume"""

        output_path = self.output_dir / f"{export_options.filename}.json"

        export_data = {
            "resume_data": resume_data,
            "export_metadata": {
                "generated_at": datetime.now().isoformat(),
                "template": export_options.template,
                "format": export_options.format,
                "ats_optimized": export_options.ats_optimized,
            },
        }

        async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
            await f.write(json.dumps(export_data, indent=2, ensure_ascii=False))

        return str(output_path)

    async def _generate_html(
        self, resume_data: Dict[str, Any], export_options: ExportOptions
    ) -> str:
        """Generate HTML format resume"""

        output_path = self.output_dir / f"{export_options.filename}.html"
        # Get template customizations
        template = self.templates.get(export_options.template)
        customizations = template.customizations if template else {}
        # Generate CSS based on template and custom styling
        css_styles = self._generate_css_styles(customizations, export_options.custom_styling)
        # Generate HTML content
        html_content = self._generate_html_content(resume_data, customizations)
        html_template = f""
        return html_template


    async def _build_docx_from_scratch(
        self, doc: Document, resume_data: Dict[str, Any], export_options: ExportOptions
    ):
        """Build DOCX document from scratch when no template is available"""

        # Personal details header
        if "personal_details" in resume_data:
            personal = resume_data["personal_details"]

            # Name as title
            name_para = doc.add_heading(personal.get("name", ""), 0)
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact info
            contact_info = []
            if personal.get("email"):
                contact_info.append(personal["email"])
            if personal.get("phone"):
                contact_info.append(personal["phone"])
            if personal.get("linkedin"):
                contact_info.append(personal["linkedin"])

            if contact_info:
                contact_para = doc.add_paragraph(" | ".join(contact_info))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Summary
            if personal.get("summary"):
                doc.add_heading("Summary", level=1)
                doc.add_paragraph(personal["summary"])

        # Work Experience
        if "work_experience" in resume_data and resume_data["work_experience"]:
            doc.add_heading("Work Experience", level=1)
            for exp in resume_data["work_experience"]:
                # Job title and company
                job_para = doc.add_paragraph()
                job_run = job_para.add_run(
                    f"{exp.get('title', '')} - {exp.get('company', '')}"
                )
                job_run.bold = True

                # Dates
                if exp.get("from_date") or exp.get("to_date"):
                    date_para = doc.add_paragraph(
                        f"{exp.get('from_date', '')} - {exp.get('to_date', 'Present')}"
                    )
                    date_para.style = "Intense Quote"

                # Summary
                if exp.get("summary"):
                    doc.add_paragraph(exp["summary"])

                # Achievements
                if exp.get("achievements"):
                    for achievement in exp["achievements"]:
                        doc.add_paragraph(achievement, style="List Bullet")

        # Education
        if "education" in resume_data and resume_data["education"]:
            doc.add_heading("Education", level=1)
            for edu in resume_data["education"]:
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(
                    f"{edu.get('degree', '')} - {edu.get('university', '')}"
                )
                edu_run.bold = True

                if edu.get("from_year") or edu.get("to_year"):
                    doc.add_paragraph(
                        f"{edu.get('from_year', '')} - {edu.get('to_year', '')}"
                    )

        # Skills
        if "skills" in resume_data and resume_data["skills"]:
            doc.add_heading("Skills", level=1)
            for skill_cat in resume_data["skills"]:
                skill_para = doc.add_paragraph()
                cat_run = skill_para.add_run(f"{skill_cat.get('category', '')}: ")
                cat_run.bold = True
                skill_para.add_run(", ".join(skill_cat.get("skills", [])))

        # Projects
        if "projects" in resume_data and resume_data["projects"]:
            doc.add_heading("Projects", level=1)
            for project in resume_data["projects"]:
                proj_para = doc.add_paragraph()
                proj_run = proj_para.add_run(project.get("name", ""))
                proj_run.bold = True

                if project.get("summary"):
                    doc.add_paragraph(project["summary"])

                if project.get("bullets"):
                    for bullet in project["bullets"]:
                        doc.add_paragraph(bullet, style="List Bullet")

    async def _batch_export_zip(self, batch_request: BatchExportRequest) -> str:
        """Export multiple resumes as a ZIP file"""

        zip_path = self.output_dir / f"{batch_request.batch_name}.zip"
        generated_files = []

        # Generate all resumes
        for i, (version, options) in enumerate(
            zip(batch_request.versions, batch_request.export_options)
        ):
            try:
                # Ensure unique filename for each version
                if (
                    not options.filename
                    or options.filename
                    == f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                ):
                    options.filename = f"{batch_request.batch_name}_version_{i+1}"

                file_path = await self.generate_resume(version, options)
                generated_files.append(file_path)
            except Exception as e:
                # Log error but continue with other files
                print(f"Error generating version {i+1}: {str(e)}")

        # Create ZIP file
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for file_path in generated_files:
                if os.path.exists(file_path):
                    zipf.write(file_path, os.path.basename(file_path))

        # Clean up individual files
        for file_path in generated_files:
            if os.path.exists(file_path):
                os.remove(file_path)

        return str(zip_path)

    async def _batch_export_individual(self, batch_request: BatchExportRequest) -> str:
        """Export multiple resumes as individual files"""

        batch_dir = self.output_dir / batch_request.batch_name
        batch_dir.mkdir(exist_ok=True)

        generated_files = []

        for i, (version, options) in enumerate(
            zip(batch_request.versions, batch_request.export_options)
        ):
            try:
                # Update output path to batch directory
                original_filename = options.filename
                options.filename = str(
                    batch_dir / f"{original_filename or f'version_{i+1}'}"
                )

                file_path = await self.generate_resume(version, options)
                generated_files.append(file_path)
            except Exception as e:
                print(f"Error generating version {i+1}: {str(e)}")

        return str(batch_dir)

    def _apply_keyword_optimizations(
        self, section_data: Any, keywords: List[str]
    ) -> Any:
        """Apply keyword optimizations to section data"""
        if isinstance(section_data, dict):
            # For work experience, add keywords to achievements
            if "achievements" in section_data and isinstance(
                section_data["achievements"], list
            ):
                for i, achievement in enumerate(section_data["achievements"]):
                    for keyword in keywords:
                        if keyword.lower() not in achievement.lower():
                            # Intelligently integrate keywords
                            section_data["achievements"][i] = self._integrate_keyword(
                                achievement, keyword
                            )

            # For skills section, add missing keywords
            if "skills" in section_data and isinstance(section_data["skills"], list):
                existing_skills = [skill.lower() for skill in section_data["skills"]]
                for keyword in keywords:
                    if keyword.lower() not in existing_skills:
                        section_data["skills"].append(keyword)

        elif isinstance(section_data, list):
            # For skills array or similar
            for section in section_data:
                if isinstance(section, dict) and "skills" in section:
                    existing_skills = [skill.lower() for skill in section["skills"]]
                    for keyword in keywords:
                        if keyword.lower() not in existing_skills:
                            section["skills"].append(keyword)

        return section_data

    def _apply_content_improvements(
        self, section_data: Any, improvements: List[Dict[str, Any]]
    ) -> Any:
        """Apply content improvements to section data"""
        if isinstance(section_data, dict):
            for improvement in improvements:
                field = improvement.get("field")
                new_content = improvement.get("content")
                improvement_type = improvement.get("type", "replace")

                if field in section_data and new_content:
                    if improvement_type == "replace":
                        section_data[field] = new_content
                    elif improvement_type == "append":
                        if isinstance(section_data[field], str):
                            section_data[field] += " " + new_content
                        elif isinstance(section_data[field], list):
                            section_data[field].append(new_content)
                    elif improvement_type == "prepend":
                        if isinstance(section_data[field], str):
                            section_data[field] = (
                                new_content + " " + section_data[field]
                            )
                        elif isinstance(section_data[field], list):
                            section_data[field].insert(0, new_content)

        return section_data

    def _apply_text_replacements(
        self, section_data: Any, replacements: List[Dict[str, str]]
    ) -> Any:
        """Apply text replacements to section data"""
        if isinstance(section_data, dict):
            for key, value in section_data.items():
                if isinstance(value, str):
                    for replacement in replacements:
                        old_text = replacement.get("old")
                        new_text = replacement.get("new")
                        if old_text and new_text and old_text in value:
                            section_data[key] = value.replace(old_text, new_text)
                elif isinstance(value, list):
                    for i, item in enumerate(value):
                        if isinstance(item, str):
                            for replacement in replacements:
                                old_text = replacement.get("old")
                                new_text = replacement.get("new")
                                if old_text and new_text and old_text in item:
                                    section_data[key][i] = item.replace(
                                        old_text, new_text
                                    )

        return section_data

    def _apply_ats_optimizations(
        self, resume_data: Dict[str, Any], ats_opts: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Apply ATS-specific optimizations"""
        optimized_data = resume_data.copy()

        # Ensure standard section names
        section_mappings = {
            "work_experience": ["experience", "employment", "work_history"],
            "education": ["academic_background", "qualifications"],
            "skills": ["technical_skills", "competencies"],
            "projects": ["portfolio", "key_projects"],
        }

        for standard_name, alternatives in section_mappings.items():
            for alt_name in alternatives:
                if alt_name in optimized_data and standard_name not in optimized_data:
                    optimized_data[standard_name] = optimized_data.pop(alt_name)

        # Remove formatting that might confuse ATS
        if ats_opts.get("remove_special_characters", False):
            optimized_data = self._remove_special_formatting(optimized_data)

        # Ensure contact information is in standard format
        if "personal_details" in optimized_data:
            personal = optimized_data["personal_details"]
            if "phone" in personal:
                personal["phone"] = self._standardize_phone_number(personal["phone"])
            if "email" in personal:
                personal["email"] = personal["email"].lower().strip()

        return optimized_data

    def _apply_job_specific_optimizations(
        self, resume_data: Dict[str, Any], job_opts: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Apply job-specific optimizations"""
        optimized_data = resume_data.copy()

        # Reorder sections based on job requirements
        if "section_priority" in job_opts:
            optimized_data = self._reorder_sections(
                optimized_data, job_opts["section_priority"]
            )

        # Emphasize relevant experience
        if "relevant_experience" in job_opts and "work_experience" in optimized_data:
            optimized_data["work_experience"] = self._prioritize_relevant_experience(
                optimized_data["work_experience"], job_opts["relevant_experience"]
            )

        # Highlight matching skills
        if "required_skills" in job_opts and "skills" in optimized_data:
            optimized_data["skills"] = self._highlight_matching_skills(
                optimized_data["skills"], job_opts["required_skills"]
            )

        return optimized_data

    def _integrate_keyword(self, text: str, keyword: str) -> str:
        """Intelligently integrate keyword into text"""
        # Simple implementation - could be enhanced with NLP
        if len(text.split()) < 10:
            return f"{text} utilizing {keyword}"
        else:
            # Insert keyword in the middle
            words = text.split()
            mid_point = len(words) // 2
            words.insert(mid_point, f"({keyword})")
            return " ".join(words)

    def _remove_special_formatting(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Remove special formatting that might confuse ATS"""
        # Remove special characters, excessive formatting, etc.
        # This is a simplified implementation
        return data

    def _standardize_phone_number(self, phone: str) -> str:
        """Standardize phone number format for ATS"""
        # Remove all non-digit characters except +
        import re

        cleaned = re.sub(r"[^\d+]", "", phone)

        # Format as +1-XXX-XXX-XXXX for US numbers
        if len(cleaned) == 10:
            return f"+1-{cleaned[:3]}-{cleaned[3:6]}-{cleaned[6:]}"
        elif len(cleaned) == 11 and cleaned.startswith("1"):
            return f"+{cleaned[0]}-{cleaned[1:4]}-{cleaned[4:7]}-{cleaned[7:]}"

        return phone  # Return original if can't standardize

    def _prioritize_relevant_experience(
        self, experiences: List[Dict[str, Any]], relevant_keywords: List[str]
    ) -> List[Dict[str, Any]]:
        """Prioritize relevant work experience"""

        def relevance_score(exp):
            score = 0
            text = f"{exp.get('title', '')} {exp.get('summary', '')} {' '.join(exp.get('achievements', []))}"
            for keyword in relevant_keywords:
                if keyword.lower() in text.lower():
                    score += 1
            return score

        return sorted(experiences, key=relevance_score, reverse=True)

    def _highlight_matching_skills(
        self, skills: List[Dict[str, Any]], required_skills: List[str]
    ) -> List[Dict[str, Any]]:
        """Highlight skills that match job requirements"""
        for skill_category in skills:
            if "skills" in skill_category:
                # Reorder skills to put matching ones first
                matching_skills = []
                other_skills = []

                for skill in skill_category["skills"]:
                    if any(
                        req_skill.lower() in skill.lower()
                        for req_skill in required_skills
                    ):
                        matching_skills.append(skill)
                    else:
                        other_skills.append(skill)

                skill_category["skills"] = matching_skills + other_skills

        return skills

    def _reorder_sections(
        self, resume_data: Dict[str, Any], section_order: List[str]
    ) -> Dict[str, Any]:
        """Reorder resume sections according to template preferences"""
        reordered = {}

        # Add sections in specified order
        for section in section_order:
            if section in resume_data:
                reordered[section] = resume_data[section]

        # Add any remaining sections
        for section, data in resume_data.items():
            if section not in reordered:
                reordered[section] = data

        return reordered


# Legacy function for backward compatibility
def generate_resume(
    parsed: dict, template: str = "modern", filetype: str = "docx"
) -> str:
    """Legacy function for backward compatibility"""
    generator = EnhancedResumeGenerator()
    export_options = ExportOptions(
        format=filetype, template=template, filename=f"resume_updated"
    )

    # Run async function in sync context
    import asyncio

    try:
        loop = asyncio.get_event_loop()
    except RuntimeError:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

    return loop.run_until_complete(generator.generate_resume(parsed, export_options))

    def _generate_css_styles(self, customizations: Dict[str, Any], custom_styling: Dict[str, Any]) -> str:
        """Generate CSS styles based on template customizations"""
        
        # Base styles
        base_styles = """
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background: white;
        }
        
        .resume-container {
            background: white;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            padding: 40px;
        }
        
        .header {
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid #333;
        }
        
        .name {
            font-size: 2.5em;
            font-weight: bold;
            margin-bottom: 10px;
            color: #2c3e50;
        }
        
        .contact-info {
            font-size: 1.1em;
            color: #666;
        }
        
        .section {
            margin: 25px 0;
        }
        
        .section-title {
            font-size: 1.4em;
            font-weight: bold;
            margin-bottom: 15px;
            padding-bottom: 5px;
            border-bottom: 1px solid #ccc;
            color: #2c3e50;
        }
        
        .experience-item, .education-item, .project-item {
            margin: 20px 0;
            padding: 15px 0;
        }
        
        .item-title {
            font-weight: bold;
            font-size: 1.1em;
            color: #2c3e50;
        }
        
        .item-subtitle {
            color: #666;
            font-style: italic;
            margin: 5px 0;
        }
        
        .item-date {
            color: #888;
            font-size: 0.9em;
        }
        
        .item-description {
            margin: 10px 0;
        }
        
        .achievements {
            list-style-type: disc;
            margin-left: 20px;
        }
        
        .achievements li {
            margin: 5px 0;
        }
        
        .skills-container {
            display: flex;
            flex-wrap: wrap;
            gap: 15px;
        }
        
        .skill-category {
            background: #f8f9fa;
            padding: 10px 15px;
            border-radius: 5px;
            border-left: 4px solid #007bff;
        }
        
        .skill-category-title {
            font-weight: bold;
            margin-bottom: 5px;
            color: #2c3e50;
        }
        
        .skill-list {
            color: #666;
        }
        
        .summary {
            font-style: italic;
            background: #f8f9fa;
            padding: 15px;
            border-left: 4px solid #007bff;
            margin: 20px 0;
        }
        """
        
        # Apply template-specific customizations
        template_styles = ""
        
        # Font family
        font_family = customizations.get("font_family", "Arial")
        template_styles += f"body {{ font-family: {font_family}, sans-serif; }}\n"
        
        # Color scheme
        color_scheme = customizations.get("color_scheme", "blue")
        color_map = {
            "blue": "#007bff",
            "green": "#28a745",
            "purple": "#6f42c1",
            "teal": "#20c997",
            "orange": "#fd7e14",
            "red": "#dc3545",
            "gray": "#6c757d",
            "black": "#000000",
            "navy": "#001f3f",
            "none": "#333333"
        }
        
        accent_color = color_map.get(color_scheme, "#007bff")
        template_styles += f"""
        .name, .section-title, .item-title {{ color: {accent_color}; }}
        .header {{ border-bottom-color: {accent_color}; }}
        .skill-category, .summary {{ border-left-color: {accent_color}; }}
        """
        
        # Spacing
        spacing = customizations.get("spacing", "standard")
        if spacing == "compact":
            template_styles += """
            .section { margin: 15px 0; }
            .experience-item, .education-item, .project-item { margin: 10px 0; padding: 10px 0; }
            .resume-container { padding: 20px; }
            """
        elif spacing == "wide":
            template_styles += """
            .section { margin: 35px 0; }
            .experience-item, .education-item, .project-item { margin: 25px 0; padding: 20px 0; }
            .resume-container { padding: 50px; }
            """
        
        # Apply custom styling overrides
        custom_styles = ""
        if custom_styling:
            if "primary_color" in custom_styling:
                custom_styles += f"""
                .name, .section-title, .item-title {{ color: {custom_styling['primary_color']}; }}
                .header {{ border-bottom-color: {custom_styling['primary_color']}; }}
                .skill-category, .summary {{ border-left-color: {custom_styling['primary_color']}; }}
                """
            
            if "font_size" in custom_styling:
                custom_styles += f"body {{ font-size: {custom_styling['font_size']}; }}\n"
            
            if "background_color" in custom_styling:
                custom_styles += f"body {{ background-color: {custom_styling['background_color']}; }}\n"
        
        return base_styles + template_styles + custom_styles

    def _generate_html_content(self, resume_data: Dict[str, Any], customizations: Dict[str, Any]) -> str:
        """Generate HTML content for resume"""
        
        content_parts = []
        
        # Header section
        if "personal_details" in resume_data:
            personal = resume_data["personal_details"]
            contact_info = []
            
            if personal.get("email"):
                contact_info.append(f'<a href="mailto:{personal["email"]}">{personal["email"]}</a>')
            if personal.get("phone"):
                contact_info.append(personal["phone"])
            if personal.get("linkedin"):
                linkedin_url = personal["linkedin"]
                if not linkedin_url.startswith("http"):
                    linkedin_url = f"https://{linkedin_url}"
                contact_info.append(f'<a href="{linkedin_url}" target="_blank">{personal["linkedin"]}</a>')
            
            header_html = f"""
            <div class="header">
                <div class="name">{personal.get('name', '')}</div>
                <div class="contact-info">{' | '.join(contact_info)}</div>
            </div>
            """
            content_parts.append(header_html)
            
            # Summary section
            if personal.get("summary"):
                summary_html = f"""
                <div class="section">
                    <div class="summary">{personal['summary']}</div>
                </div>
                """
                content_parts.append(summary_html)
        
        # Get section order from template
        section_order = customizations.get("section_order", [
            "work_experience", "education", "skills", "projects"
        ])
        
        # Generate sections in specified order
        for section_name in section_order:
            if section_name in resume_data and resume_data[section_name]:
                section_html = self._generate_section_html(section_name, resume_data[section_name])
                if section_html:
                    content_parts.append(section_html)
        
        return "\n".join(content_parts)

    def _generate_section_html(self, section_name: str, section_data: Any) -> str:
        """Generate HTML for a specific section"""
        
        section_titles = {
            "work_experience": "Work Experience",
            "education": "Education",
            "skills": "Skills",
            "projects": "Projects",
            "certifications": "Certifications",
            "awards": "Awards"
        }
        
        title = section_titles.get(section_name, section_name.replace("_", " ").title())
        
        if section_name == "work_experience":
            return self._generate_work_experience_html(title, section_data)
        elif section_name == "education":
            return self._generate_education_html(title, section_data)
        elif section_name == "skills":
            return self._generate_skills_html(title, section_data)
        elif section_name == "projects":
            return self._generate_projects_html(title, section_data)
        else:
            return ""

    def _generate_work_experience_html(self, title: str, experiences: List[Dict[str, Any]]) -> str:
        """Generate HTML for work experience section"""
        
        items_html = []
        for exp in experiences:
            achievements_html = ""
            if exp.get("achievements"):
                achievements_list = "\n".join([f"<li>{achievement}</li>" for achievement in exp["achievements"]])
                achievements_html = f'<ul class="achievements">{achievements_list}</ul>'
            
            date_range = ""
            if exp.get("from_date") or exp.get("to_date"):
                date_range = f'<div class="item-date">{exp.get("from_date", "")} - {exp.get("to_date", "Present")}</div>'
            
            summary_html = ""
            if exp.get("summary"):
                summary_html = f'<div class="item-description">{exp["summary"]}</div>'
            
            item_html = f"""
            <div class="experience-item">
                <div class="item-title">{exp.get('title', '')}</div>
                <div class="item-subtitle">{exp.get('company', '')}</div>
                {date_range}
                {summary_html}
                {achievements_html}
            </div>
            """
            items_html.append(item_html)
        
        return f"""
        <div class="section">
            <div class="section-title">{title}</div>
            {"".join(items_html)}
        </div>
        """

    def _generate_education_html(self, title: str, education: List[Dict[str, Any]]) -> str:
        """Generate HTML for education section"""
        
        items_html = []
        for edu in education:
            date_range = ""
            if edu.get("from_year") or edu.get("to_year"):
                date_range = f'<div class="item-date">{edu.get("from_year", "")} - {edu.get("to_year", "")}</div>'
            
            item_html = f"""
            <div class="education-item">
                <div class="item-title">{edu.get('degree', '')}</div>
                <div class="item-subtitle">{edu.get('university', '')}</div>
                {date_range}
            </div>
            """
            items_html.append(item_html)
        
        return f"""
        <div class="section">
            <div class="section-title">{title}</div>
            {"".join(items_html)}
        </div>
        """

    def _generate_skills_html(self, title: str, skills: List[Dict[str, Any]]) -> str:
        """Generate HTML for skills section"""
        
        skills_html = []
        for skill_cat in skills:
            category_name = skill_cat.get("category", "")
            skill_list = ", ".join(skill_cat.get("skills", []))
            
            skill_html = f"""
            <div class="skill-category">
                <div class="skill-category-title">{category_name}</div>
                <div class="skill-list">{skill_list}</div>
            </div>
            """
            skills_html.append(skill_html)
        
        return f"""
        <div class="section">
            <div class="section-title">{title}</div>
            <div class="skills-container">
                {"".join(skills_html)}
            </div>
        </div>
        """

    def _generate_projects_html(self, title: str, projects: List[Dict[str, Any]]) -> str:
        """Generate HTML for projects section"""
        
        items_html = []
        for project in projects:
            bullets_html = ""
            if project.get("bullets"):
                bullets_list = "\n".join([f"<li>{bullet}</li>" for bullet in project["bullets"]])
                bullets_html = f'<ul class="achievements">{bullets_list}</ul>'
            
            summary_html = ""
            if project.get("summary"):
                summary_html = f'<div class="item-description">{project["summary"]}</div>'
            
            item_html = f"""
            <div class="project-item">
                <div class="item-title">{project.get('name', '')}</div>
                {summary_html}
                {bullets_html}
            </div>
            """
            items_html.append(item_html)
        
        return f"""
        <div class="section">
            <div class="section-title">{title}</div>
            {"".join(items_html)}
        </div>
        """