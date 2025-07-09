from docxtpl import DocxTemplate
from docx import Document
import os
from jinja2 import Template


def generate_resume(
    parsed: dict, template: str = "modern", filetype: str = "docx"
) -> str:
    # For demo, use a simple docx template
    template_path = os.path.join("templates", f"{template}.docx")
    output_path = os.path.join("generated", f"resume_updated.{filetype}")
    if not os.path.exists(template_path):
        # fallback: create a simple docx
        doc = Document()
        doc.add_heading(parsed.get("personal_details", {}).get("name", "Name"), 0)
        doc.add_paragraph(
            "Contact: " + parsed.get("personal_details", {}).get("contact", "")
        )
        doc.add_paragraph(
            "LinkedIn: " + parsed.get("personal_details", {}).get("linkedin", "")
        )
        doc.add_heading("Education", level=1)
        doc.add_paragraph(parsed.get("education", ""))
        doc.add_heading("Work Experience", level=1)
        doc.add_paragraph(parsed.get("work_experience", ""))
        doc.add_heading("Projects", level=1)
        doc.add_paragraph(parsed.get("projects", ""))
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(parsed.get("skills", ""))
        doc.save(output_path)
    else:
        doc = DocxTemplate(template_path)
        doc.render(parsed)
        doc.save(output_path)
    # PDF export can be added with docx2pdf or similar
    return output_path
